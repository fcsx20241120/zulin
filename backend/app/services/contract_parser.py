"""
合同模板导入服务
"""

from docx import Document
import re
from datetime import datetime
from typing import Dict, Any


def parse_contract_template(file_path: str) -> Dict[str, Any]:
    """
    解析租赁合同 docx 模板

    Args:
        file_path: docx 文件路径

    Returns:
        解析后的合同数据字典
    """
    doc = Document(file_path)

    # 提取所有文本
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    result = {"landlord": {}, "tenant": {}, "house": {}, "lease": {}}

    # 解析房东信息
    landlord_match = re.search(r"出租方 \(简称甲方\)：\s*(.+?)(?:\n|$)", text)
    if landlord_match:
        result["landlord"]["name"] = landlord_match.group(1).strip()

    # 解析租客信息
    tenant_match = re.search(r"承租方 \(简称乙方\)：\s*(.+?)(?:\n|$)", text)
    if tenant_match:
        result["tenant"]["name"] = tenant_match.group(1).strip()

    # 解析租赁期限
    years_match = re.search(r"共 (\d+) 年", text)
    if years_match:
        result["lease"]["lease_years"] = int(years_match.group(1))

    # 解析租金
    rent_match = re.search(r"每月租金为 (\d+) 元", text)
    if rent_match:
        result["lease"]["monthly_rent"] = float(rent_match.group(1))

    # 解析支付方式
    payment_match = re.search(r"按 (\w+) 支付", text)
    if payment_match:
        result["lease"]["payment_type"] = payment_match.group(1)

    # 解析保证金
    deposit_match = re.search(r"保证金 (\d+) 元", text)
    if deposit_match:
        result["lease"]["deposit"] = float(deposit_match.group(1))

    return result


def parse_contract_from_bytes(file_bytes: bytes) -> Dict[str, Any]:
    """
    从字节流解析合同

    Args:
        file_bytes: docx 文件字节流

    Returns:
        解析后的合同数据字典
    """
    from io import BytesIO

    doc = Document(BytesIO(file_bytes))

    # 提取所有文本
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    result = {"landlord": {}, "tenant": {}, "house": {}, "lease": {}, "furniture": []}

    # 简单解析关键信息
    patterns = {
        "landlord_name": r"出租方.*?：\s*(.+?)(?:\n|$)",
        "tenant_name": r"承租方.*?：\s*(.+?)(?:\n|$)",
        "lease_years": r"共 (\d+) 年",
        "monthly_rent": r"(\d+) 元.*?月租",
        "deposit": r"保证金.*?(\d+) 元",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            value = match.group(1)
            if key in ["lease_years"]:
                value = int(value)
            elif key in ["monthly_rent", "deposit"]:
                value = float(value)

            if key.startswith("landlord_"):
                result["landlord"][key.replace("landlord_", "")] = value
            elif key.startswith("tenant_"):
                result["tenant"][key.replace("tenant_", "")] = value
            else:
                result["lease"][key] = value

    return result
